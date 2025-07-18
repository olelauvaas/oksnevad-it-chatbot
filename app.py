import streamlit as st
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_chroma import Chroma
from langchain_core.messages import SystemMessage, HumanMessage
import os

# === KONFIG ===
API_NØKKEL = "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"
FILNAVN = "faq.docx"
CHROMA_DB_DIR = "chroma_db"

# === Leser Word-dokument ===
@st.cache_resource
def bygg_vector_db():
    # Henter tekst fra Word
    doc = Document(FILNAVN)
    tekst = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    # Deler opp i biter
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    dokumenter = [LCDocument(page_content=chunk) for chunk in splitter.split_text(tekst)]

    # Lager og returnerer vektordatabase
    embedding_model = OpenAIEmbeddings(openai_api_key=API_NØKKEL)
    db = Chroma.from_documents(dokumenter, embedding_model, persist_directory=CHROMA_DB_DIR)
    return db

# === Funksjon for å spørre boten ===
def spør_bot(spørsmål, db):
    chat = ChatOpenAI(model="gpt-4o", openai_api_key=API_NØKKEL, temperature=0.3)

    # Hent dokumenter med score
    docs_med_score = db.similarity_search_with_score(spørsmål, k=3)
    relevante_docs = [doc for doc, score in docs_med_score if score < 0.85]

    systemmelding = SystemMessage(content="""\
Du er en pålitelig og hjelpsom digital assistent ved en videregående skole. Du svarer nøkternt, kort og tydelig – uten å dikte opp informasjon.
Gi korte svar hvis det er nok, men utvid ved behov, spesielt ved spørsmål som krever forklaring eller bakgrunn.
Unngå humor og pynt, og hold deg til fakta.
Svar gjerne med punktliste hvis det gjør svaret tydeligere.
""")

    if relevante_docs:
        innhold = "\n".join([doc.page_content for doc in relevante_docs])
        meldinger = [
            systemmelding,
            HumanMessage(content=f"Denne informasjonen kan være relevant:\n{innhold}\n\nSpørsmål: {spørsmål}")
        ]
    else:
        meldinger = [
            systemmelding,
            HumanMessage(content=spørsmål)
        ]

    svar = chat.invoke(meldinger)
    return svar.content

# === Streamlit-grensesnitt ===
st.set_page_config(page_title="IT-kontoret på Øksnevad vgs", page_icon="💻")
st.image("logo.png", width=200)
st.title("💻 IT-kontoret på Øksnevad vgs")
st.write("Stil et spørsmål basert på dokumentet og få svar med en gang.")

# Bygg databasen (hvis ikke allerede gjort)
db = bygg_vector_db()

# Spørreskjema med vanlige spørsmål
st.subheader("🔎 Vanlige spørsmål")

# Liste med vanlige spørsmål
spørsmål_valg = [
    "Hva er PC-ordningen?",
    "Hvor laster jeg ned Office?",
    "Hvordan kobler jeg til printer?",
    "Hvordan logger jeg inn på Teams?",
    "Hva gjør jeg hvis PC-en er ødelagt?"
]

# Bruk session_state for å lagre valgt spørsmål
if "valgt_spørsmål" not in st.session_state:
    st.session_state.valgt_spørsmål = ""

# Vis knapper
for spm in spørsmål_valg:
    if st.button(spm):
        st.session_state.valgt_spørsmål = spm

# Manuelt spørsmål (brukeren kan endre eller skrive nytt)
spørsmål = st.text_input(
    "Eller skriv inn et annet spørsmål:",
    value=st.session_state.valgt_spørsmål,
    placeholder="Eks: Hvordan får jeg tilgang til e-post?"
)

# Hvis det finnes et spørsmål – vis svar
if spørsmål:
    with st.spinner("🤔 Tenker litt..."):
        svar = spør_bot(spørsmål, db)
    st.success("💬 Svar fra bot:")
    st.write(svar)